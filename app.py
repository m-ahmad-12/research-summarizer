import streamlit as st
import groq
import PyPDF2
import io
from docx import Document
from fpdf import FPDF

API_KEY = "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"

def get_client():
    return groq.Groq(api_key=API_KEY)

def extract_pdf_text(uploaded_file):
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def call_groq(client, prompt):
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1024,
    )
    return response.choices[0].message.content

def summarize_section(client, text, section):
    prompt = f"""From the research paper below, extract and summarize only the {section} section.
Be concise and clear. If the section is not found, write 'Not found in paper.'
Paper:
{text[:6000]}
Summary of {section}:"""
    return call_groq(client, prompt)

def rewrite_section(client, text):
    prompt = f"""Rewrite the following text in a completely original way.
Keep the meaning intact but change the wording, sentence structure, and style.
Make it plagiarism-free and academic in tone.
Text:
{text}
Rewritten:"""
    return call_groq(client, prompt)

def generate_citations(client, text):
    prompt = f"""From the research paper below, extract all references or citations mentioned.
Format each one in APA style. If no references found, generate 3 likely academic citations based on the paper topic.
Paper:
{text[:6000]}
Citations (APA format):"""
    return call_groq(client, prompt)

def generate_layman_summary(client, text):
    prompt = f"""Read the research paper below and explain it in very simple English.
Pretend you are explaining it to a 15-year-old with no technical background.
Use short sentences, simple words, and relatable examples.
Paper:
{text[:6000]}
Simple Explanation:"""
    return call_groq(client, prompt)

def extract_keywords(client, text):
    prompt = f"""From the research paper below, extract the top 10 most important research keywords.
Return them as a numbered list only. No explanation needed.
Paper:
{text[:6000]}
Top 10 Keywords:"""
    return call_groq(client, prompt)

def suggest_future_work(client, text):
    prompt = f"""Based on the research paper below, suggest 5 future research directions or improvements.
Be specific and academic. Format as a numbered list.
Paper:
{text[:6000]}
Future Work Suggestions:"""
    return call_groq(client, prompt)

def answer_question(client, text, question):
    prompt = f"""You are an expert research assistant. Answer the following question based only on the research paper provided.
Be accurate, concise, and cite specific parts of the paper if possible.
If the answer is not in the paper, say 'This information is not available in the paper.'
Paper:
{text[:6000]}
Question: {question}
Answer:"""
    return call_groq(client, prompt)

def export_word(data):
    doc = Document()
    doc.add_heading("Research Paper Summary", 0)
    for section, content in data.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def export_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.set_font("Arial", style="B", size=16)
    pdf.cell(200, 10, txt="Research Paper Summary", ln=True, align="C")
    pdf.ln(5)
    for section, content in data.items():
        pdf.set_font("Arial", style="B", size=13)
        pdf.cell(200, 10, txt=section, ln=True)
        pdf.set_font("Arial", size=11)
        clean = content.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 8, txt=clean)
        pdf.ln(3)
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf

# ── PAGE CONFIG ───────────────────────────────────────────────────────────────
st.set_page_config(page_title="Research Summarizer", page_icon="📄", layout="wide")

st.markdown("""
<style>
    .stApp { background-color: #f0f4f8; }

    .main-title {
        text-align: center;
        font-size: 2.4rem;
        font-weight: 800;
        color: #1a1a2e !important;
        padding: 20px 0 5px 0;
    }

    .main-subtitle {
        text-align: center;
        font-size: 1rem;
        color: #444444 !important;
        margin-bottom: 30px;
    }

    .results-heading {
        font-size: 1.6rem;
        font-weight: 800;
        color: #1a1a2e !important;
        margin-bottom: 16px;
    }

    .section-card {
        background-color: #ffffff;
        border-radius: 16px;
        padding: 20px 24px 10px 24px;
        margin-bottom: 8px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.07);
        border-left: 6px solid;
    }

    .section-title {
        font-size: 1.25rem;
        font-weight: 700;
    }

    .col-label {
        font-size: 0.82rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-bottom: 8px;
        padding: 4px 12px;
        border-radius: 20px;
        display: inline-block;
    }

    .label-summary  { background-color: #dbeafe; color: #1e40af !important; }
    .label-rewrite  { background-color: #dcfce7; color: #166534 !important; }
    .label-special  { background-color: #fef3c7; color: #92400e !important; }

    .content-box {
        background-color: #f1f5f9;
        border-radius: 10px;
        padding: 16px;
        font-size: 0.95rem;
        line-height: 1.75;
        color: #1e293b !important;
        min-height: 120px;
        border: 1px solid #e2e8f0;
    }

    .feature-card {
        background-color: #ffffff;
        border-radius: 16px;
        padding: 20px 24px;
        margin-bottom: 16px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.07);
        border-left: 6px solid;
    }

    .citations-card {
        background-color: #ffffff;
        border-radius: 16px;
        padding: 20px 24px;
        margin: 16px 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.07);
        border-left: 6px solid #f5a623;
    }

    .citations-title {
        font-size: 1.25rem;
        font-weight: 700;
        color: #b45309 !important;
        margin-bottom: 12px;
    }

    .chat-bubble-user {
        background-color: #1a73e8;
        color: white !important;
        padding: 10px 16px;
        border-radius: 16px 16px 4px 16px;
        margin: 8px 0;
        max-width: 80%;
        margin-left: auto;
        font-size: 0.95rem;
    }

    .chat-bubble-ai {
        background-color: #ffffff;
        color: #1e293b !important;
        padding: 10px 16px;
        border-radius: 16px 16px 16px 4px;
        margin: 8px 0;
        max-width: 80%;
        border: 1px solid #e2e8f0;
        font-size: 0.95rem;
    }

    .keyword-badge {
        display: inline-block;
        background-color: #ede9fe;
        color: #5b21b6 !important;
        padding: 4px 12px;
        border-radius: 20px;
        margin: 4px;
        font-size: 0.88rem;
        font-weight: 600;
    }

    .section-divider {
        border: none;
        border-top: 2px dashed #cbd5e1;
        margin: 18px 0;
    }
</style>
""", unsafe_allow_html=True)

section_styles = {
    "Abstract":     {"color": "#6c63ff", "icon": "🔍"},
    "Introduction": {"color": "#1a73e8", "icon": "📖"},
    "Methodology":  {"color": "#db2777", "icon": "🔬"},
    "Results":      {"color": "#059669", "icon": "📊"},
    "Conclusion":   {"color": "#d97706", "icon": "🏁"},
}

# ── HEADER ────────────────────────────────────────────────────────────────────
st.markdown('<div class="main-title">📄 Research Paper Summarizer</div>', unsafe_allow_html=True)
st.markdown('<div class="main-subtitle">Upload a research paper PDF to get section-wise summaries, plagiarism-free rewrites, citations, keywords, and more.</div>', unsafe_allow_html=True)

# ── SIDEBAR ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## ⚙️ Controls")
    st.markdown("---")
    uploaded_file = st.file_uploader("📂 Upload Research Paper (PDF)", type=["pdf"])
    st.markdown("---")
    analyze_btn = st.button("🚀 Analyze Paper", use_container_width=True)
    st.markdown("---")
    st.markdown("**How to use:**")
    st.markdown("1. Upload your PDF\n2. Click Analyze\n3. View results\n4. Chat with paper\n5. Download report")

# ── SESSION STATE ─────────────────────────────────────────────────────────────
sections = ["Abstract", "Introduction", "Methodology", "Results", "Conclusion"]

for key in ["results", "rewrites", "citations", "layman", "keywords", "future_work", "chat_history", "paper_text"]:
    if key not in st.session_state:
        st.session_state[key] = {} if key in ["results", "rewrites"] else ([] if key == "chat_history" else "")

# ── ANALYZE ───────────────────────────────────────────────────────────────────
if analyze_btn:
    if not uploaded_file:
        st.error("⚠️ Please upload a PDF file first.")
    else:
        client = get_client()
        with st.spinner("📖 Reading PDF..."):
            text = extract_pdf_text(uploaded_file)
            st.session_state.paper_text = text

        st.session_state.results      = {}
        st.session_state.rewrites     = {}
        st.session_state.chat_history = []

        total = len(sections) * 2 + 4
        step  = 0
        progress = st.progress(0, text="Starting analysis...")

        for section in sections:
            progress.progress(step / total, text=f"Summarizing {section}...")
            st.session_state.results[section] = summarize_section(client, text, section)
            step += 1
            progress.progress(step / total, text=f"Rewriting {section}...")
            st.session_state.rewrites[section] = rewrite_section(client, st.session_state.results[section])
            step += 1

        progress.progress(step / total, text="Generating citations...")
        st.session_state.citations = generate_citations(client, text)
        step += 1

        progress.progress(step / total, text="Writing layman summary...")
        st.session_state.layman = generate_layman_summary(client, text)
        step += 1

        progress.progress(step / total, text="Extracting keywords...")
        st.session_state.keywords = extract_keywords(client, text)
        step += 1

        progress.progress(step / total, text="Suggesting future work...")
        st.session_state.future_work = suggest_future_work(client, text)
        progress.progress(1.0, text="✅ Done!")

# ── RESULTS ───────────────────────────────────────────────────────────────────
if st.session_state.results:
    st.markdown("---")
    st.markdown('<div class="results-heading">📋 Analysis Results</div>', unsafe_allow_html=True)

    for section in sections:
        style = section_styles[section]
        color = style["color"]
        icon  = style["icon"]

        st.markdown(f"""
        <div class="section-card" style="border-left-color:{color};">
            <div class="section-title" style="color:{color};">{icon}&nbsp; {section}</div>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        summary_text = st.session_state.results.get(section, "")
        rewrite_text = st.session_state.rewrites.get(section, "")

        with col1:
            st.markdown('<span class="col-label label-summary">📝 Summary</span>', unsafe_allow_html=True)
            st.markdown(f'<div class="content-box">{summary_text}</div>', unsafe_allow_html=True)
            st.code(summary_text, language=None)

        with col2:
            st.markdown('<span class="col-label label-rewrite">✅ Plagiarism-Free Rewrite</span>', unsafe_allow_html=True)
            st.markdown(f'<div class="content-box">{rewrite_text}</div>', unsafe_allow_html=True)
            st.code(rewrite_text, language=None)

        st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    # ── LAYMAN SUMMARY ────────────────────────────────────────────────────────
    st.markdown("""
    <div class="feature-card" style="border-left-color:#0ea5e9;">
        <div class="section-title" style="color:#0369a1;">🧒 Layman Summary (Simple English)</div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown(f'<div class="content-box">{st.session_state.layman}</div>', unsafe_allow_html=True)
    st.code(st.session_state.layman, language=None)
    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    # ── KEYWORDS ──────────────────────────────────────────────────────────────
    st.markdown("""
    <div class="feature-card" style="border-left-color:#7c3aed;">
        <div class="section-title" style="color:#5b21b6;">🏷️ Top Research Keywords</div>
    </div>
    """, unsafe_allow_html=True)
    keywords_list = [k.strip().lstrip("0123456789. ") for k in st.session_state.keywords.split("\n") if k.strip()]
    keyword_html = " ".join([f'<span class="keyword-badge">{kw}</span>' for kw in keywords_list if kw])
    st.markdown(f'<div class="content-box" style="min-height:60px;">{keyword_html}</div>', unsafe_allow_html=True)
    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    # ── FUTURE WORK ───────────────────────────────────────────────────────────
    st.markdown("""
    <div class="feature-card" style="border-left-color:#10b981;">
        <div class="section-title" style="color:#065f46;">🔮 Future Work Suggestions</div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown(f'<div class="content-box">{st.session_state.future_work.replace(chr(10), "<br>")}</div>', unsafe_allow_html=True)
    st.code(st.session_state.future_work, language=None)
    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    # ── CITATIONS ─────────────────────────────────────────────────────────────
    st.markdown(f"""
    <div class="citations-card">
        <div class="citations-title">📚 Citations (APA Format)</div>
        <div class="content-box">{st.session_state.citations.replace(chr(10), '<br>')}</div>
    </div>
    """, unsafe_allow_html=True)
    st.code(st.session_state.citations, language=None)
    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    # ── Q&A CHAT ──────────────────────────────────────────────────────────────
    st.markdown("""
    <div class="feature-card" style="border-left-color:#f43f5e;">
        <div class="section-title" style="color:#be123c;">💬 Chat with Your Paper</div>
    </div>
    """, unsafe_allow_html=True)

    for chat in st.session_state.chat_history:
        st.markdown(f'<div class="chat-bubble-user">🧑 {chat["question"]}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="chat-bubble-ai">🤖 {chat["answer"]}</div>', unsafe_allow_html=True)

    with st.form("chat_form", clear_on_submit=True):
        user_q = st.text_input("Ask anything about the paper...", placeholder="e.g. What dataset was used?")
        submit = st.form_submit_button("Send 💬")

    if submit and user_q.strip():
        client = get_client()
        with st.spinner("Thinking..."):
            ans = answer_question(client, st.session_state.paper_text, user_q)
        st.session_state.chat_history.append({"question": user_q, "answer": ans})
        st.rerun()

    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    # ── DOWNLOADS ─────────────────────────────────────────────────────────────
    st.markdown('<div class="results-heading">⬇️ Download Report</div>', unsafe_allow_html=True)

    export_data = {}
    for section in sections:
        export_data[f"{section} - Summary"] = st.session_state.results.get(section, "")
        export_data[f"{section} - Rewrite"] = st.session_state.rewrites.get(section, "")
    export_data["Layman Summary"]   = st.session_state.layman
    export_data["Keywords"]         = st.session_state.keywords
    export_data["Future Work"]      = st.session_state.future_work
    export_data["Citations"]        = st.session_state.citations

    col1, col2 = st.columns(2)
    with col1:
        word_file = export_word(export_data)
        st.download_button("📝 Download Word Document", word_file,
                           file_name="research_summary.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    with col2:
        pdf_file = export_pdf(export_data)
        st.download_button("📄 Download PDF Report", pdf_file,
                           file_name="research_summary.pdf",
                           mime="application/pdf",
                           use_container_width=True)
